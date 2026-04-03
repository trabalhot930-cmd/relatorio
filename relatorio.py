import streamlit as st
from PIL import Image as PILImage
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import io
import copy
from datetime import date
import os
import streamlit.components.v1 as components

# =========================
# LOGIN
# =========================
usuarios = ["Juan", "Bruno", "Josiel"]
senha_padrao = "BM123"

if "logado" not in st.session_state:
    st.session_state.logado = False

if not st.session_state.logado:
    st.title("🔐 Acesso Técnico")
    usuario = st.selectbox("Selecione seu nome", usuarios)
    senha = st.text_input("Senha", type="password")

    if st.button("Entrar"):
        if senha == senha_padrao:
            st.session_state.logado = True
            st.session_state.usuario = usuario
            st.success("✅ Acesso liberado!")
            st.rerun()
        else:
            st.error("❌ Senha incorreta")
    st.stop()

# =========================
# CONFIG
# =========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_PATH = os.path.join(BASE_DIR, "modelo.docx")

MESES = {
    1:"Janeiro", 2:"Fevereiro", 3:"Março", 4:"Abril", 5:"Maio", 6:"Junho",
    7:"Julho", 8:"Agosto", 9:"Setembro", 10:"Outubro", 11:"Novembro", 12:"Dezembro"
}

# =========================
# FUNÇÕES AUXILIARES
# =========================
def full_text(p):
    return "".join(r.text for r in p.runs)

def substituir_paragrafo(p, novo_texto):
    """Substitui texto preservando formatação do primeiro run."""
    if not p.runs:
        return
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rPr = p.runs[0]._r.find(f'{{{ns}}}rPr')
    rPr_copia = copy.deepcopy(rPr) if rPr is not None else None

    for run in list(p.runs):
        run._r.getparent().remove(run._r)

    novo_run = p.add_run(novo_texto)
    if rPr_copia is not None:
        novo_run._r.insert(0, rPr_copia)

def substituir_bloco(doc, p_inicio_idx, p_fim_idx, linhas):
    """
    Remove parágrafos de p_inicio_idx até p_fim_idx (inclusive)
    e insere novos parágrafos com as linhas fornecidas.
    """
    paragrafos = doc.paragraphs
    p_ref = paragrafos[p_inicio_idx]
    parent = p_ref._element.getparent()
    idx_insert = list(parent).index(p_ref._element)

    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    rPr_ref = None
    if p_ref.runs:
        rPr_ref = p_ref.runs[0]._r.find(f'{{{ns}}}rPr')

    # Remove parágrafos do bloco (de trás para frente)
    for i in range(p_fim_idx, p_inicio_idx - 1, -1):
        elem = doc.paragraphs[i]._element
        elem.getparent().remove(elem)

    # Insere novos parágrafos
    for i, linha in enumerate(linhas):
        novo_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(f'{{{ns}}}val', 'ListParagraph')
        pPr.append(pStyle)
        novo_p.append(pPr)

        r = OxmlElement('w:r')
        if rPr_ref is not None:
            r.append(copy.deepcopy(rPr_ref))
        t = OxmlElement('w:t')
        t.text = linha
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        novo_p.append(r)
        parent.insert(idx_insert + i, novo_p)

# =========================
# INTERFACE
# =========================
st.title("📄 Relatório Norte Energia")
st.success(f"👷 Técnico logado: {st.session_state.usuario}")

numero     = st.text_input("Número do Relatório", placeholder="Ex: 1")
assunto    = st.text_input("Assunto", placeholder="Ex: MANUTENÇÃO RADAR CANAL DE FUGA")
data_manut = st.date_input("Data", value=date.today())
local      = st.text_input("Localidade", placeholder="Ex: Canal de Fuga")

# =========================
# PARTE INFORMATIVA (editável)
# =========================
st.subheader("📝 Parte Informativa")
texto_informativo = st.text_area(
    "Escreva o texto informativo:",
    value=(
        f"Manutenção Radar {local or '[local]'}\n"
        f"\n"
        f"A equipe de Meios Eletrônicos, sob a Superintendência de Segurança Corporativa, "
        f"executou na data de {data_manut.day} de {MESES[data_manut.month].lower()} "
        f"de {data_manut.year} a manutenção do sistema radar da localidade {local or '[local]'}.\n"
        f"\n"
        f"Foram executadas as atividades de:\n"
        f"• Testes\n"
        f"• Religamento (equipamento estava congelado)"
    ),
    height=250
)

# =========================
# PARTE ILUSTRATIVA (editável)
# =========================
st.subheader("🖼️ Parte Ilustrativa")
texto_ilustrativo = st.text_area(
    "Nome da imagem (aparece acima da foto em azul):",
    value=f"Manutenção {local or '[local]'}",
    height=80
)

# =========================
# PARTE CONCLUSIVA (editável)
# =========================
st.subheader("📌 Parte Conclusiva")
texto_conclusivo = st.text_area(
    "Escreva o texto conclusivo:",
    value="Após as manutenções os equipamentos foram recolocados em operação.",
    height=120
)

# =========================
# GPS
# =========================
st.subheader("📍 Localização automática")
gps_html = """
<script>
navigator.geolocation.getCurrentPosition(
    (position) => {
        const lat = position.coords.latitude;
        const lon = position.coords.longitude;
        const coords = lat + "," + lon;
        window.parent.postMessage({type: "streamlit:setComponentValue", value: coords}, "*");
    }
);
</script>
"""
coords = components.html(gps_html, height=0)
if coords:
    st.success(f"📍 GPS: {coords}")

# =========================
# FOTO
# =========================
st.subheader("📸 Foto da Atividade")
opcao_foto = st.radio("Escolha:", ["Tirar foto", "Enviar da galeria"])

foto_bytes = None
if opcao_foto == "Tirar foto":
    foto = st.camera_input("Abrir câmera")
    if foto:
        foto_bytes = foto.getvalue()
else:
    arquivo = st.file_uploader("Selecionar imagem", type=["jpg", "jpeg", "png"])
    if arquivo:
        foto_bytes = arquivo.read()

# =========================
# GERAR RELATÓRIO
# =========================
if st.button("🚀 Gerar Relatório"):

    if not os.path.exists(TEMPLATE_PATH):
        st.error("❌ modelo.docx não encontrado!")
        st.stop()

    # Validação dos campos obrigatórios
    erros = []
    if not numero.strip():            erros.append("Número do Relatório")
    if not assunto.strip():           erros.append("Assunto")
    if not local.strip():             erros.append("Localidade")
    if not texto_informativo.strip(): erros.append("Parte Informativa")
    if not texto_ilustrativo.strip(): erros.append("Parte Ilustrativa")
    if not texto_conclusivo.strip():  erros.append("Parte Conclusiva")

    if erros:
        st.error(f"❌ Preencha os campos obrigatórios: {', '.join(erros)}")
        st.stop()

    doc = Document(TEMPLATE_PATH)

    data_str   = f"{data_manut.day} de {MESES[data_manut.month]} de {data_manut.year}"
    data_upper = data_str.upper()

    # --------------------------------------------------
    # 1) CABEÇALHO — dentro da tabela
    # --------------------------------------------------
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = full_text(p)
                    if "RELATÓRIO DE SEGURANÇA" in txt:
                        substituir_paragrafo(p, f"RELATÓRIO DE SEGURANÇA Nr. {numero} / 2026")
                    elif "DATA:" in txt:
                        substituir_paragrafo(p, f"DATA: {data_upper}")
                    elif "ASSUNTO:" in txt:
                        substituir_paragrafo(p, f"ASSUNTO: {assunto.upper()}")

    # --------------------------------------------------
    # 2) PARTE INFORMATIVA
    #    Parágrafos 3 a 7 no template original
    # --------------------------------------------------
    # Encontra início e fim do bloco informativo
    idx_info_ini = None
    idx_info_fim = None
    for i, p in enumerate(doc.paragraphs):
        txt = full_text(p)
        if idx_info_ini is None and txt.strip() == 'Manutenção Radar canal de fuga':
            idx_info_ini = i
        if idx_info_ini is not None and 'PARTE ILUSTRATIVA' in txt:
            # O bloco termina no parágrafo vazio logo antes desta seção
            idx_info_fim = i - 2
            break

    if idx_info_ini is not None and idx_info_fim is not None and idx_info_ini <= idx_info_fim:
        substituir_bloco(doc, idx_info_ini, idx_info_fim, texto_informativo.split('\n'))

    # --------------------------------------------------
    # 3) PARTE ILUSTRATIVA — caption ACIMA da foto, azul (estilo Legenda)
    # --------------------------------------------------
    for i, p in enumerate(doc.paragraphs):
        if 'PARTE ILUSTRATIVA' in full_text(p):
            if i + 1 < len(doc.paragraphs):
                p_caption = doc.paragraphs[i + 1]
                # Remove todos os runs mantendo o estilo "Legenda" (azul, bold, centralizado)
                for run in list(p_caption.runs):
                    run._r.getparent().remove(run._r)
                # Adiciona novo run — a cor azul vem do estilo Legenda automaticamente
                novo_run = p_caption.add_run(texto_ilustrativo)
            break

    # --------------------------------------------------
    # 4) FOTO — cabe na primeira página, ajustada ao espaço disponível
    # --------------------------------------------------
    if foto_bytes:
        for i, p in enumerate(doc.paragraphs):
            if 'PARTE ILUSTRATIVA' in full_text(p):
                try:
                    # i+1 = caption (nome da imagem, já preenchido acima)
                    # i+2 = parágrafo para a foto
                    p_foto = doc.paragraphs[i + 2]
                    p_foto.clear()
                    run = p_foto.add_run()
                    # Tamanho fixo definido pelo cliente: 6,8cm largura x 12,09cm altura
                    run.add_picture(
                        io.BytesIO(foto_bytes),
                        width=Cm(6.8),
                        height=Cm(12.09)
                    )
                    p_foto.alignment = WD_ALIGN_PARAGRAPH.CENTER

                except Exception as e:
                    st.error(f"Erro ao inserir imagem: {e}")
                break

    # --------------------------------------------------
    # 5) PARTE CONCLUSIVA
    # --------------------------------------------------
    idx_conc_ini = None
    idx_conc_fim = None
    for i, p in enumerate(doc.paragraphs):
        txt = full_text(p)
        if 'PARTE CONCLUSIVA' in txt and idx_conc_ini is None:
            idx_conc_ini = i + 1
        if 'Vitória do Xingu' in txt and idx_conc_ini is not None:
            # Acha último parágrafo não-vazio antes de Vitória do Xingu
            for j in range(i - 1, idx_conc_ini - 1, -1):
                if full_text(doc.paragraphs[j]).strip():
                    idx_conc_fim = j
                    break
            if idx_conc_fim is None:
                idx_conc_fim = i - 1
            break

    if idx_conc_ini is not None and idx_conc_fim is not None and idx_conc_ini <= idx_conc_fim:
        substituir_bloco(doc, idx_conc_ini, idx_conc_fim, texto_conclusivo.split('\n'))

    # --------------------------------------------------
    # 6) DATA FINAL — Vitória do Xingu
    # --------------------------------------------------
    for p in doc.paragraphs:
        if 'Vitória do Xingu' in full_text(p):
            substituir_paragrafo(p, f"Vitória do Xingu /PA, {data_str}")
            break

    # --------------------------------------------------
    # 7) EXPORTAR
    # --------------------------------------------------
    buffer = io.BytesIO()
    doc.save(buffer)

    st.success("✅ Relatório gerado com sucesso!")
    st.download_button(
        "📥 Baixar Relatório",
        buffer.getvalue(),
        f"Relatorio_{numero}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # WhatsApp
    mensagem = f"Relatório Nr {numero}\nAssunto: {assunto}\nLocal: {local}"
